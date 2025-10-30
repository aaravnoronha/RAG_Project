"""
RAG Agent System - A framework for building AI agents on top of language models
with Retrieval-Augmented Generation capabilities.

This system provides:
- Document ingestion and vector storage
- Semantic search and retrieval
- Agent orchestration with tool usage
- Memory management
- Conversation handling
"""

import os
import json
import asyncio
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from enum import Enum
from datetime import datetime
import hashlib
from abc import ABC, abstractmethod

# Vector store and embedding dependencies
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
import tiktoken

# For LLM interaction
from openai import OpenAI
import anthropic

# Document processing
from pypdf import PdfReader
import docx
import markdown
from bs4 import BeautifulSoup


# ============= Core Data Models =============

@dataclass
class Document:
    """Represents a document in the knowledge base"""
    id: str
    content: str
    metadata: Dict[str, Any]
    embedding: Optional[np.ndarray] = None
    chunks: List['DocumentChunk'] = field(default_factory=list)
    
    def __post_init__(self):
        if not self.id:
            self.id = hashlib.md5(self.content.encode()).hexdigest()


@dataclass
class DocumentChunk:
    """Represents a chunk of a document for retrieval"""
    id: str
    document_id: str
    content: str
    metadata: Dict[str, Any]
    embedding: Optional[np.ndarray] = None
    chunk_index: int
    
    def __post_init__(self):
        if not self.id:
            self.id = f"{self.document_id}_{self.chunk_index}"


@dataclass
class Message:
    """Represents a message in a conversation"""
    role: str  # 'user', 'assistant', 'system'
    content: str
    timestamp: datetime = field(default_factory=datetime.now)
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class AgentMemory:
    """Manages agent's short-term and long-term memory"""
    short_term: List[Message] = field(default_factory=list)
    long_term: List[Message] = field(default_factory=list)
    context_window: int = 10
    
    def add_message(self, message: Message):
        """Add message to memory"""
        self.short_term.append(message)
        if len(self.short_term) > self.context_window:
            self.long_term.append(self.short_term.pop(0))
    
    def get_context(self) -> List[Message]:
        """Get current context for the agent"""
        return self.short_term[-self.context_window:]
    
    def search_memory(self, query: str, limit: int = 5) -> List[Message]:
        """Search through all memory for relevant messages"""
        # Simple keyword search - could be enhanced with semantic search
        all_messages = self.short_term + self.long_term
        relevant = []
        query_lower = query.lower()
        
        for msg in all_messages:
            if query_lower in msg.content.lower():
                relevant.append(msg)
                if len(relevant) >= limit:
                    break
        
        return relevant


# ============= Document Processing =============

class DocumentProcessor:
    """Handles document ingestion and chunking"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.tokenizer = tiktoken.get_encoding("cl100k_base")
    
    def process_document(self, file_path: str, metadata: Dict[str, Any] = None) -> Document:
        """Process a document file and return Document object"""
        metadata = metadata or {}
        content = self._extract_text(file_path)
        
        doc = Document(
            id=hashlib.md5(f"{file_path}_{content[:100]}".encode()).hexdigest(),
            content=content,
            metadata={
                **metadata,
                'source': file_path,
                'processed_at': datetime.now().isoformat()
            }
        )
        
        # Create chunks
        doc.chunks = self._chunk_text(doc)
        return doc
    
    def _extract_text(self, file_path: str) -> str:
        """Extract text from various file formats"""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self._extract_pdf(file_path)
        elif ext == '.docx':
            return self._extract_docx(file_path)
        elif ext == '.md':
            return self._extract_markdown(file_path)
        elif ext in ['.txt', '.py', '.js', '.json', '.yaml', '.yml']:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    
    def _extract_markdown(self, file_path: str) -> str:
        """Extract text from Markdown"""
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        html = markdown.markdown(md_content)
        soup = BeautifulSoup(html, 'html.parser')
        return soup.get_text()
    
    def _chunk_text(self, document: Document) -> List[DocumentChunk]:
        """Split document into overlapping chunks"""
        tokens = self.tokenizer.encode(document.content)
        chunks = []
        
        for i in range(0, len(tokens), self.chunk_size - self.chunk_overlap):
            chunk_tokens = tokens[i:i + self.chunk_size]
            chunk_text = self.tokenizer.decode(chunk_tokens)
            
            chunk = DocumentChunk(
                id=f"{document.id}_{len(chunks)}",
                document_id=document.id,
                content=chunk_text,
                metadata={
                    **document.metadata,
                    'chunk_index': len(chunks),
                    'token_count': len(chunk_tokens)
                },
                chunk_index=len(chunks)
            )
            chunks.append(chunk)
        
        return chunks


# ============= Vector Store =============

class VectorStore:
    """Manages embeddings and similarity search using FAISS"""
    
    def __init__(self, embedding_model: str = "all-MiniLM-L6-v2", dimension: int = 384):
        self.embedding_model = SentenceTransformer(embedding_model)
        self.dimension = dimension
        self.index = faiss.IndexFlatL2(dimension)
        self.documents: Dict[str, DocumentChunk] = {}
        self.id_to_index: Dict[str, int] = {}
    
    def add_documents(self, chunks: List[DocumentChunk]):
        """Add document chunks to the vector store"""
        texts = [chunk.content for chunk in chunks]
        embeddings = self.embedding_model.encode(texts)
        
        for chunk, embedding in zip(chunks, embeddings):
            chunk.embedding = embedding
            idx = len(self.documents)
            self.documents[chunk.id] = chunk
            self.id_to_index[chunk.id] = idx
            self.index.add(embedding.reshape(1, -1))
    
    def search(self, query: str, k: int = 5) -> List[Tuple[DocumentChunk, float]]:
        """Search for similar documents"""
        query_embedding = self.embedding_model.encode([query])
        distances, indices = self.index.search(query_embedding, k)
        
        results = []
        for i, (idx, distance) in enumerate(zip(indices[0], distances[0])):
            if idx < len(self.documents):
                chunk_id = list(self.documents.keys())[idx]
                chunk = self.documents[chunk_id]
                results.append((chunk, float(distance)))
        
        return results
    
    def save(self, path: str):
        """Save vector store to disk"""
        faiss.write_index(self.index, f"{path}.index")
        with open(f"{path}.meta", 'w') as f:
            meta = {
                'documents': {k: {
                    'id': v.id,
                    'document_id': v.document_id,
                    'content': v.content,
                    'metadata': v.metadata,
                    'chunk_index': v.chunk_index
                } for k, v in self.documents.items()},
                'id_to_index': self.id_to_index
            }
            json.dump(meta, f)
    
    def load(self, path: str):
        """Load vector store from disk"""
        self.index = faiss.read_index(f"{path}.index")
        with open(f"{path}.meta", 'r') as f:
            meta = json.load(f)
            self.documents = {
                k: DocumentChunk(**v) for k, v in meta['documents'].items()
            }
            self.id_to_index = meta['id_to_index']


# ============= Tools and Actions =============

class Tool(ABC):
    """Base class for agent tools"""
    
    @property
    @abstractmethod
    def name(self) -> str:
        pass
    
    @property
    @abstractmethod
    def description(self) -> str:
        pass
    
    @abstractmethod
    async def execute(self, *args, **kwargs) -> Any:
        pass


class WebSearchTool(Tool):
    """Tool for web searching"""
    
    @property
    def name(self) -> str:
        return "web_search"
    
    @property
    def description(self) -> str:
        return "Search the web for current information"
    
    async def execute(self, query: str) -> List[Dict[str, str]]:
        # This would integrate with a real search API
        # For demo, returning mock results
        return [
            {"title": "Result 1", "snippet": f"Information about {query}", "url": "https://example.com/1"},
            {"title": "Result 2", "snippet": f"More about {query}", "url": "https://example.com/2"}
        ]


class CalculatorTool(Tool):
    """Tool for mathematical calculations"""
    
    @property
    def name(self) -> str:
        return "calculator"
    
    @property
    def description(self) -> str:
        return "Perform mathematical calculations"
    
    async def execute(self, expression: str) -> float:
        try:
            # Safe evaluation of mathematical expressions
            result = eval(expression, {"__builtins__": {}}, {})
            return result
        except Exception as e:
            return f"Error: {str(e)}"


class CodeExecutorTool(Tool):
    """Tool for executing Python code"""
    
    @property
    def name(self) -> str:
        return "code_executor"
    
    @property
    def description(self) -> str:
        return "Execute Python code snippets"
    
    async def execute(self, code: str) -> str:
        try:
            # In production, use a sandboxed environment
            exec_globals = {}
            exec(code, exec_globals)
            return str(exec_globals.get('result', 'Code executed successfully'))
        except Exception as e:
            return f"Error: {str(e)}"


# ============= LLM Providers =============

class LLMProvider(ABC):
    """Base class for LLM providers"""
    
    @abstractmethod
    async def generate(self, messages: List[Message], **kwargs) -> str:
        pass
    
    @abstractmethod
    async def generate_with_tools(self, messages: List[Message], tools: List[Tool], **kwargs) -> Tuple[str, Optional[Dict]]:
        pass


class OpenAIProvider(LLMProvider):
    """OpenAI LLM provider"""
    
    def __init__(self, api_key: str, model: str = "gpt-4"):
        self.client = OpenAI(api_key=api_key)
        self.model = model
    
    async def generate(self, messages: List[Message], **kwargs) -> str:
        response = self.client.chat.completions.create(
            model=self.model,
            messages=[{"role": m.role, "content": m.content} for m in messages],
            **kwargs
        )
        return response.choices[0].message.content
    
    async def generate_with_tools(self, messages: List[Message], tools: List[Tool], **kwargs) -> Tuple[str, Optional[Dict]]:
        # Convert tools to OpenAI function format
        functions = [
            {
                "name": tool.name,
                "description": tool.description,
                "parameters": {
                    "type": "object",
                    "properties": {},
                    "required": []
                }
            }
            for tool in tools
        ]
        
        response = self.client.chat.completions.create(
            model=self.model,
            messages=[{"role": m.role, "content": m.content} for m in messages],
            functions=functions,
            function_call="auto",
            **kwargs
        )
        
        message = response.choices[0].message
        
        if message.function_call:
            return message.content or "", {
                "name": message.function_call.name,
                "arguments": json.loads(message.function_call.arguments)
            }
        
        return message.content, None


class AnthropicProvider(LLMProvider):
    """Anthropic Claude provider"""
    
    def __init__(self, api_key: str, model: str = "claude-3-opus-20240229"):
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = model
    
    async def generate(self, messages: List[Message], **kwargs) -> str:
        # Convert messages to Anthropic format
        system_message = next((m.content for m in messages if m.role == "system"), "")
        user_messages = [{"role": m.role, "content": m.content} 
                        for m in messages if m.role != "system"]
        
        response = self.client.messages.create(
            model=self.model,
            system=system_message,
            messages=user_messages,
            **kwargs
        )
        return response.content[0].text
    
    async def generate_with_tools(self, messages: List[Message], tools: List[Tool], **kwargs) -> Tuple[str, Optional[Dict]]:
        # Anthropic doesn't have native function calling like OpenAI
        # We'll implement a prompt-based approach
        tool_descriptions = "\n".join([
            f"- {tool.name}: {tool.description}"
            for tool in tools
        ])
        
        system_prompt = f"""You have access to these tools:
{tool_descriptions}

To use a tool, respond with:
TOOL: tool_name
ARGS: {{"arg1": "value1", "arg2": "value2"}}

Otherwise, respond normally."""
        
        messages_with_system = [Message(role="system", content=system_prompt)] + messages
        response = await self.generate(messages_with_system, **kwargs)
        
        # Parse for tool usage
        if "TOOL:" in response and "ARGS:" in response:
            lines = response.split("\n")
            tool_name = None
            args = None
            
            for i, line in enumerate(lines):
                if line.startswith("TOOL:"):
                    tool_name = line.replace("TOOL:", "").strip()
                elif line.startswith("ARGS:"):
                    args_str = line.replace("ARGS:", "").strip()
                    try:
                        args = json.loads(args_str)
                    except:
                        args = {}
            
            if tool_name:
                return "", {"name": tool_name, "arguments": args or {}}
        
        return response, None


# ============= RAG Agent =============

class RAGAgent:
    """Main RAG Agent class that orchestrates everything"""
    
    def __init__(
        self,
        llm_provider: LLMProvider,
        vector_store: VectorStore,
        tools: List[Tool] = None,
        system_prompt: str = None
    ):
        self.llm = llm_provider
        self.vector_store = vector_store
        self.tools = tools or []
        self.memory = AgentMemory()
        self.system_prompt = system_prompt or self._default_system_prompt()
        
        # Add tools to dictionary for easy access
        self.tool_map = {tool.name: tool for tool in self.tools}
    
    def _default_system_prompt(self) -> str:
        return """You are a helpful AI assistant with access to a knowledge base and various tools.
        Use the retrieved context to answer questions accurately.
        If you need current information or to perform calculations, use the available tools.
        Always cite your sources when using retrieved information."""
    
    async def process_query(self, query: str, use_rag: bool = True, use_tools: bool = True) -> str:
        """Process a user query and generate a response"""
        
        # Add user message to memory
        user_message = Message(role="user", content=query)
        self.memory.add_message(user_message)
        
        # Retrieve relevant context if RAG is enabled
        context = ""
        if use_rag and self.vector_store:
            results = self.vector_store.search(query, k=5)
            if results:
                context = "\n\n".join([
                    f"[Source {i+1}]: {chunk.content}"
                    for i, (chunk, _) in enumerate(results)
                ])
        
        # Prepare messages for LLM
        messages = self._prepare_messages(query, context)
        
        # Generate response with or without tools
        if use_tools and self.tools:
            response, tool_call = await self.llm.generate_with_tools(messages, self.tools)
            
            # Execute tool if requested
            if tool_call:
                tool_result = await self._execute_tool(tool_call)
                
                # Add tool result to context and regenerate
                tool_message = Message(
                    role="system",
                    content=f"Tool '{tool_call['name']}' returned: {tool_result}"
                )
                messages.append(tool_message)
                response = await self.llm.generate(messages)
        else:
            response = await self.llm.generate(messages)
        
        # Add assistant response to memory
        assistant_message = Message(role="assistant", content=response)
        self.memory.add_message(assistant_message)
        
        return response
    
    def _prepare_messages(self, query: str, context: str) -> List[Message]:
        """Prepare messages for LLM"""
        messages = [Message(role="system", content=self.system_prompt)]
        
        if context:
            messages.append(Message(
                role="system",
                content=f"Retrieved context:\n{context}"
            ))
        
        # Add conversation history
        messages.extend(self.memory.get_context())
        
        return messages
    
    async def _execute_tool(self, tool_call: Dict) -> Any:
        """Execute a tool call"""
        tool_name = tool_call['name']
        args = tool_call.get('arguments', {})
        
        if tool_name in self.tool_map:
            tool = self.tool_map[tool_name]
            return await tool.execute(**args)
        else:
            return f"Tool '{tool_name}' not found"
    
    def add_documents(self, file_paths: List[str], processor: DocumentProcessor = None):
        """Add documents to the knowledge base"""
        if not processor:
            processor = DocumentProcessor()
        
        for file_path in file_paths:
            print(f"Processing {file_path}...")
            doc = processor.process_document(file_path)
            self.vector_store.add_documents(doc.chunks)
            print(f"Added {len(doc.chunks)} chunks from {file_path}")
    
    async def chat(self):
        """Interactive chat interface"""
        print("RAG Agent initialized. Type 'exit' to quit.")
        print("-" * 50)
        
        while True:
            user_input = input("\nYou: ").strip()
            
            if user_input.lower() == 'exit':
                print("Goodbye!")
                break
            
            response = await self.process_query(user_input)
            print(f"\nAssistant: {response}")


# ============= Agent Factory =============

class AgentFactory:
    """Factory for creating different types of agents"""
    
    @staticmethod
    def create_research_agent(api_key: str, provider: str = "openai") -> RAGAgent:
        """Create a research-focused agent"""
        llm = OpenAIProvider(api_key) if provider == "openai" else AnthropicProvider(api_key)
        vector_store = VectorStore()
        tools = [WebSearchTool(), CalculatorTool()]
        
        system_prompt = """You are a research assistant specialized in finding and synthesizing information.
        Always cite your sources and provide comprehensive answers based on the available data."""
        
        return RAGAgent(llm, vector_store, tools, system_prompt)
    
    @staticmethod
    def create_coding_agent(api_key: str, provider: str = "openai") -> RAGAgent:
        """Create a coding-focused agent"""
        llm = OpenAIProvider(api_key) if provider == "openai" else AnthropicProvider(api_key)
        vector_store = VectorStore()
        tools = [CodeExecutorTool(), WebSearchTool()]
        
        system_prompt = """You are a coding assistant specialized in helping with programming tasks.
        You can write, review, and execute code. Always explain your code and provide best practices."""
        
        return RAGAgent(llm, vector_store, tools, system_prompt)
    
    @staticmethod
    def create_qa_agent(api_key: str, provider: str = "openai") -> RAGAgent:
        """Create a Q&A agent for document-based questions"""
        llm = OpenAIProvider(api_key) if provider == "openai" else AnthropicProvider(api_key)
        vector_store = VectorStore()
        
        system_prompt = """You are a Q&A assistant that answers questions based on provided documents.
        Always base your answers on the retrieved context and cite the relevant sources."""
        
        return RAGAgent(llm, vector_store, [], system_prompt)


# ============= Example Usage =============

async def main():
    """Example usage of the RAG Agent system"""
    
    # Initialize components
    print("Initializing RAG Agent System...")
    
    # Use environment variable for API key (you'll need to set this)
    api_key = os.getenv("OPENAI_API_KEY", "your-api-key-here")
    
    # Create a research agent
    agent = AgentFactory.create_research_agent(api_key)
    
    # Example: Add documents to knowledge base
    # processor = DocumentProcessor()
    # agent.add_documents(["path/to/document1.pdf", "path/to/document2.md"])
    
    # Example: Process a query
    query = "What are the key principles of RAG systems?"
    response = await agent.process_query(query)
    print(f"Query: {query}")
    print(f"Response: {response}")
    
    # Start interactive chat
    # await agent.chat()


if __name__ == "__main__":
    # Run the example
    asyncio.run(main())
